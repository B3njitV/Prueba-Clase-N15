from docx import Document

def ingreso_datos():
    nombre = input("Ingrese el nombre del trabajador: ").strip()[:30]
    sueldo_base = float(input("Ingrese el sueldo base del trabajador: "))
    horas_extras = float(input("Ingrese el número de horas extras trabajadas en el mes: "))
    return nombre, sueldo_base, horas_extras

def calcular_liquidacion(sueldo_base, horas_extras):
    pago_horas_extras = horas_extras * (sueldo_base / 160) * 1.5
    total_ingresos = sueldo_base + pago_horas_extras
    descuento_fonasa = total_ingresos * 0.07
    descuento_afp = total_ingresos * 0.1
    sueldo_neto = total_ingresos - descuento_fonasa - descuento_afp
    return pago_horas_extras, total_ingresos, descuento_fonasa, descuento_afp, sueldo_neto

def mostrar_liquidacion(nombre, sueldo_base, pago_horas_extras, total_ingresos, descuento_fonasa, descuento_afp, sueldo_neto):
    print("\n--- Liquidación de Sueldo ---")
    print(f"Nombre del trabajador: {nombre}")
    print(f"Sueldo base: ${sueldo_base:.2f}")
    print(f"Pago por horas extras: ${pago_horas_extras:.2f}")
    print(f"Total de ingresos: ${total_ingresos:.2f}")
    print(f"Descuento por FONASA: ${descuento_fonasa:.2f}")
    print(f"Descuento por AFP: ${descuento_afp:.2f}")
    print(f"Sueldo neto a pagar: ${sueldo_neto:.2f}")

def generar_archivo(nombre, sueldo_base, pago_horas_extras, total_ingresos, descuento_fonasa, descuento_afp, sueldo_neto):
    document = Document()
    document.add_heading('Liquidación de Sueldo', level=1)
    document.add_paragraph(f"Nombre del trabajador: {nombre}")
    document.add_paragraph(f"Sueldo base: ${sueldo_base:.2f}")
    document.add_paragraph(f"Pago por horas extras: ${pago_horas_extras:.2f}")
    document.add_paragraph(f"Total de ingresos: ${total_ingresos:.2f}")
    document.add_paragraph(f"Descuento por FONASA: ${descuento_fonasa:.2f}")
    document.add_paragraph(f"Descuento por AFP: ${descuento_afp:.2f}")
    document.add_paragraph(f"Sueldo neto a pagar: ${sueldo_neto:.2f}")
    file_name = f"liquidacion_{nombre}.docx"
    document.save(file_name)
    print(f"Archivo de liquidación generado: {file_name}")

def main():
    while True:
        nombre, sueldo_base, horas_extras = ingreso_datos()
        pago_horas_extras, total_ingresos, descuento_fonasa, descuento_afp, sueldo_neto = calcular_liquidacion(sueldo_base, horas_extras)
        mostrar_liquidacion(nombre, sueldo_base, pago_horas_extras, total_ingresos, descuento_fonasa, descuento_afp, sueldo_neto)
        generar_archivo(nombre, sueldo_base, pago_horas_extras, total_ingresos, descuento_fonasa, descuento_afp, sueldo_neto)
        if input("\n¿Desea calcular la liquidación de otro trabajador? (s/n): ").lower() != 's':
            break

if __name__ == "__main__":
    main()
